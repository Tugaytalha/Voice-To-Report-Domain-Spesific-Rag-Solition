import os
import subprocess
from docx import Document
from query_data import query_rag, PROMPT_TEMPLATE
from populate_database import clear_database
from get_embedding_function import get_embedding_function
from langchain_community.llms.ollama import Ollama

EVAL_PROMPT = """
Expected Response: {expected_response}
Actual Response: {actual_response}
---
(Answer with 'true' or 'false') Does the actual response match the expected response? 
"""

TEST_QUERIES = {

}

EMBEDDING_MODELS = [
    # "emrecan/bert-base-turkish-cased-mean-nli-stsb-tr",
    # "emrecan/convbert-base-turkish-mc4-cased-allnli_tr",
    "atasoglu/roberta-small-turkish-clean-uncased-nli-stsb-tr",
    "atasoglu/distilbert-base-turkish-cased-nli-stsb-tr",
    "atasoglu/xlm-roberta-base-nli-stsb-tr",
    "atasoglu/mbert-base-cased-nli-stsb-tr",
    "Omerhan/intfloat-fine-tuned-14376-v4",
    # "atasoglu/turkish-base-bert-uncased-mean-nli-stsb-tr",
    "jinaai/jina-embeddings-v3",
]

def test_rag_with_embeddings(embedding_model_name):
    """
    Tests the RAG application with a specific embedding model.
    """
    # Create a new DOCX document for the report
    document = Document()
    document.add_heading("RAG Application Test Report", 0)

    print(f"Testing with embedding: {embedding_model_name}")
    document.add_heading(f"Embedding: {embedding_model_name}", level=1)

    env = os.environ.copy()

    # Set current venv path as working venv path
    env['PATH'] = os.path.join(os.getcwd(), 'venv', 'bin') + ':' + env['PATH']


    # Initialize embedding function with appropriate settings
    command = ["./venv/scripts/python", "populate_database.py", "--reset", "--model-type", "sentence-transformer", "--model-name", embedding_model_name]
    subprocess.run(command, check=True, env=env)

    # Test with each query
    for query, expected_response in TEST_QUERIES.items():
        try:
            response, retrieved_chunks = query_rag(query, get_embedding_function(embedding_model_name, True))
        except Exception as e:
            print(f"Error during query processing: {e}")
            response = "Error during query processing"
            retrieved_chunks = []

        evaluation_result = evaluate_response(response, expected_response)

        # Add results to the document
        document.add_paragraph(f"Query: {query}")
        document.add_paragraph(f"Expected Response: {expected_response}")
        document.add_paragraph(f"Actual Response: {response}")
        document.add_paragraph(f"Evaluation: {evaluation_result}")

        # Add retrieved chunks to the document
        document.add_paragraph("Retrieved Chunks:")
        for chunk in retrieved_chunks:
            document.add_paragraph(f" - Source: {chunk['source']}, Content: {chunk['content']}, Score: {chunk['score']}")

        document.add_paragraph("---")

    # Save the document

    document.save((f"rag_test_report_{embedding_model_name}.docx").replace("/", "_"))
    print(f"RAG test report generated: rag_test_report_{embedding_model_name}.docx")

def evaluate_response(actual_response, expected_response):
    """
    Evaluates the actual response against the expected response using an LLM.
    """
    model = Ollama(model="llama3.1:8b")
    prompt = EVAL_PROMPT.format(
        expected_response=expected_response, actual_response=actual_response
    )
    evaluation_result = model.invoke(prompt)
    return evaluation_result.strip()

def main():
    """
    Runs the RAG tests for all embedding models.
    """
    for embedding_model_name in EMBEDDING_MODELS:
        test_rag_with_embeddings(embedding_model_name)

if __name__ == "__main__":
    main()